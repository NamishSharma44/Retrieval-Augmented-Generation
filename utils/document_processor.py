"""
Advanced Document Processing Module with Enhanced Error Handling
"""
from langchain_community.document_loaders import (
    PyPDFLoader, 
    Docx2txtLoader,
    TextLoader
)
from langchain_text_splitters import RecursiveCharacterTextSplitter
from typing import List, Optional
from pathlib import Path
import logging
import os

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Advanced document processor with multiple loader fallbacks"""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        
        # Initialize text splitter with smart separators
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""],
            keep_separator=True
        )
        
        logger.info(f"DocumentProcessor initialized: chunk_size={chunk_size}, overlap={chunk_overlap}")
    
    def _load_pdf(self, file_path: str) -> List:
        """Load PDF with enhanced error handling"""
        logger.info(f"Loading PDF: {Path(file_path).name}")
        
        try:
            loader = PyPDFLoader(file_path)
            documents = loader.load()
            
            if not documents:
                logger.warning(f"No content extracted from PDF: {file_path}")
                return []
            
            # Validate content
            total_chars = sum(len(doc.page_content) for doc in documents)
            logger.info(f"✅ PDF loaded: {len(documents)} pages, {total_chars} characters")
            
            # Filter empty pages
            documents = [doc for doc in documents if doc.page_content.strip()]
            
            return documents
            
        except Exception as e:
            logger.error(f"❌ Error loading PDF {file_path}: {str(e)}")
            return []
    
    def _load_docx(self, file_path: str) -> List:
        """Load DOCX with multiple fallback methods"""
        logger.info(f"Loading Word document: {Path(file_path).name}")
        
        # Try primary loader (Docx2txtLoader)
        try:
            loader = Docx2txtLoader(file_path)
            documents = loader.load()
            
            if documents and documents[0].page_content.strip():
                total_chars = len(documents[0].page_content)
                logger.info(f"✅ DOCX loaded with Docx2txtLoader: {total_chars} characters")
                return documents
        except Exception as e:
            logger.warning(f"Docx2txtLoader failed: {str(e)}, trying fallback...")
        
        # Fallback: Manual extraction with python-docx
        try:
            import docx
            doc = docx.Document(file_path)
            
            # Extract all text
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        full_text.append(row_text)
            
            content = "\n\n".join(full_text)
            
            if content.strip():
                from langchain_core.documents import Document
                documents = [Document(
                    page_content=content,
                    metadata={"source": Path(file_path).name}
                )]
                logger.info(f"✅ DOCX loaded with python-docx: {len(content)} characters")
                return documents
                
        except Exception as e:
            logger.error(f"❌ All DOCX loaders failed for {file_path}: {str(e)}")
        
        return []
    
    def _load_txt(self, file_path: str) -> List:
        """Load plain text files"""
        logger.info(f"Loading text file: {Path(file_path).name}")
        
        try:
            loader = TextLoader(file_path, encoding='utf-8')
            documents = loader.load()
            
            if documents:
                total_chars = len(documents[0].page_content)
                logger.info(f"✅ Text file loaded: {total_chars} characters")
                return documents
        except UnicodeDecodeError:
            # Try different encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    loader = TextLoader(file_path, encoding=encoding)
                    documents = loader.load()
                    logger.info(f"✅ Text file loaded with {encoding} encoding")
                    return documents
                except:
                    continue
        except Exception as e:
            logger.error(f"❌ Error loading text file {file_path}: {str(e)}")
        
        return []
    
    def load_documents(self, file_paths: List[str]) -> List:
        """Load multiple documents with enhanced processing"""
        logger.info(f"\n📄 Loading {len(file_paths)} file(s)...")
        all_documents = []
        
        for file_path in file_paths:
            file_name = Path(file_path).name
            logger.info(f"\n--- Processing: {file_name} ---")
            
            # Validate file
            if not os.path.exists(file_path):
                logger.warning(f"⚠️ File not found: {file_path}")
                continue
            
            file_size = os.path.getsize(file_path)
            logger.info(f"File size: {file_size:,} bytes ({file_size/1024/1024:.2f} MB)")
            
            if file_size == 0:
                logger.warning(f"⚠️ File is empty: {file_path}")
                continue
            
            # Load based on extension
            try:
                file_ext = Path(file_path).suffix.lower()
                
                if file_ext == '.pdf':
                    docs = self._load_pdf(file_path)
                elif file_ext in ['.docx', '.doc']:
                    docs = self._load_docx(file_path)
                elif file_ext == '.txt':
                    docs = self._load_txt(file_path)
                else:
                    logger.error(f"❌ Unsupported file format: {file_ext}")
                    continue
                
                if not docs:
                    logger.warning(f"⚠️ No content extracted from {file_name}")
                    continue
                
                # Add enhanced metadata
                for doc in docs:
                    doc.metadata.update({
                        'source': file_name,
                        'file_path': file_path,
                        'file_type': file_ext,
                        'file_size': file_size
                    })
                    
                    # Validate content
                    if not doc.page_content or not doc.page_content.strip():
                        logger.warning(f"⚠️ Empty content in {file_name}")
                
                all_documents.extend(docs)
                logger.info(f"✅ Successfully added {len(docs)} document(s)")
                
            except Exception as e:
                logger.error(f"❌ Error processing {file_name}: {str(e)}")
                import traceback
                traceback.print_exc()
        
        logger.info(f"\n📊 Total documents loaded: {len(all_documents)}")
        return all_documents
    
    def chunk_documents(self, documents: List) -> List:
        """Chunk documents with validation"""
        if not documents:
            logger.error("❌ No documents to chunk!")
            return []
        
        logger.info(f"\n✂️ Chunking {len(documents)} document(s)...")
        
        # Filter valid documents
        valid_docs = [
            doc for doc in documents 
            if doc.page_content and doc.page_content.strip()
        ]
        
        logger.info(f"Valid documents with content: {len(valid_docs)}")
        
        if not valid_docs:
            logger.error("❌ All documents are empty!")
            return []
        
        try:
            chunks = self.text_splitter.split_documents(valid_docs)
            logger.info(f"✅ Created {len(chunks)} chunks")
            
            if chunks:
                # Statistics
                avg_length = sum(len(c.page_content) for c in chunks) / len(chunks)
                logger.info(f"📊 Average chunk length: {avg_length:.0f} characters")
                
                # Sample
                sample = chunks[0]
                logger.info(f"\n📄 Sample chunk:")
                logger.info(f"   Length: {len(sample.page_content)} characters")
                logger.info(f"   Source: {sample.metadata.get('source', 'Unknown')}")
                logger.info(f"   Preview: {sample.page_content[:200]}...")
            
            return chunks
            
        except Exception as e:
            logger.error(f"❌ Error chunking documents: {str(e)}")
            import traceback
            traceback.print_exc()
            return []
    
    def process_files(self, file_paths: List[str]) -> List:
        """Complete processing pipeline"""
        documents = self.load_documents(file_paths)
        if not documents:
            raise ValueError("No documents could be loaded")
        
        chunks = self.chunk_documents(documents)
        if not chunks:
            raise ValueError("No chunks created from documents")
        
        return chunks